#pip3 install python-docx
import sys
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt,Inches,RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

textfile = sys.argv[1]

with open(textfile) as f1:
    textarray = f1.readlines()

m = re.search( "(.*).txt", textfile)

document = Document()
table = document.add_table(1, 3)
table.style = "TableGrid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

i=1

for text in textarray:
	text = text.strip()
	cells = table.add_row().cells
	cells[0].paragraphs[0].add_run("%s"%(i))
	cells[1].paragraphs[0].add_run(text.split("\t")[0])
	cells[2].paragraphs[0].add_run(text.split("\t")[1])
	#cells[3].paragraphs[0].add_run(text.split("\t")[2])
	i=i+1
#randomString = ''.join(random.choice(string.ascii_uppercase + string.digits) for m in range(10))
fileName = m.group(1)+'.docx'
print (fileName)
#fileUrl  = publicFilePath+fileName
#print fileUrl
#print filePath
document.save(fileName)
