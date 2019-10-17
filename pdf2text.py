
import sys
import subprocess
import re
#from docx import Document
import os
from os import listdir
from os.path import isfile, join


pdfdir = sys.argv[1]
files = [f for f in os.listdir(pdfdir)]# if os.path.isfile(f)]
for f in files:
	#pdfinput = os.path.abspath(f)
	pdfinput = pdfdir + "/" + f
	#print(pdfinput)
	pdfinput = re.sub(r' ',r'\ ', pdfinput, flags=re.MULTILINE)
	command = "pdftotext "+ pdfinput #+ " -"
	#print(command)
	subprocess.getoutput(command)#.encode('utf8')
	#break
"""
pdfcontent = str(pdfcontent)

#pdfcontent = pdfcontent.encode('UTF-8')
document = Document()
#pdfcontent = re.sub(r'\r\n', '', pdfcontent)
#pdfcontent = re.sub(r'\\n',r'<br>', pdfcontent, flags=re.MULTILINE)
p = document.add_paragraph(pdfcontent)
pdfinput = re.sub(r'\\ ',r' ', pdfinput, flags=re.MULTILINE)

filename = re.sub(r'\.pdf$','.doc',pdfinput)
document.save(filename)
"""

#print(stdout)
"""
import PyPDF2
import sys

pdfFileObj = open(sys.argv[1],'rb')     #'rb' for read binary mode
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
total_pages = pdfReader.numPages

page_content = ''
for i in range(1,total_pages):
	pageObj = pdfReader.getPage(i)          #'9' is the page number
	page_content += pageObj.extractText()

print(page_content)

"""
"""
import sys
from tika import parser
from docx import Document
import re
import os
from os import listdir
from os.path import isfile, join


files = [f for f in os.listdir('/home/nagaraju/python-git/utilities/pdftest')]# if os.path.isfile(f)]
print(files)
for f in files:
	f = os.path.abspath(f)
	f = re.sub(r' ',r'\ ', f, flags=re.MULTILINE)
	document = Document()
	print(f)
	pdfinput = f
	raw = parser.from_file(f)
	#print(raw['content'])
	pdfcontent = raw['content']
	pdfcontent = re.sub(r'\n\n', '\n', pdfcontent, flags=re.MULTILINE)
	#pdfcontent = re.sub(r'\r\n', '', pdfcontent)

	p = document.add_paragraph(pdfcontent)
	filename = re.sub(r'\.pdf$','.doc',pdfinput)
	document.save(filename)
"""