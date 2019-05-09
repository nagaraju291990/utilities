#input a docx file consisting of table this script will convert it into a tab seperated text to standard output
from docx import Document
import sys

document = Document(sys.argv[1])
tables = document.tables
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text,end='')
            print("\t",end='')
        print()